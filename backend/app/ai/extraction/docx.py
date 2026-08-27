"""DOCX document extraction."""
from typing import List, Tuple

from docx import Document


class DOCXExtractor:
    """Extract text from DOCX documents."""

    @staticmethod
    def extract(file_path: str) -> Tuple[str, List[dict]]:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file.

        Returns:
            Tuple of:
            - Full extracted text
            - List of paragraph metadata with text

        Raises:
            Exception: If DOCX cannot be read or has no extractable text.
        """
        try:
            doc = Document(file_path)
            
            if len(doc.paragraphs) == 0:
                raise Exception("DOCX contains no paragraphs")
            
            full_text = []
            paragraphs_metadata = []
            
            for para_num, paragraph in enumerate(doc.paragraphs, start=1):
                text = paragraph.text.strip()
                
                if text:
                    full_text.append(text)
                    paragraphs_metadata.append({
                        "section": para_num,
                        "text": text
                    })
            
            # Also extract table text
            for table_num, table in enumerate(doc.tables, start=1):
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    combined_table_text = "\n".join(table_text)
                    full_text.append(combined_table_text)
                    paragraphs_metadata.append({
                        "section": f"Table_{table_num}",
                        "text": combined_table_text
                    })
            
            if not full_text:
                raise Exception("DOCX contains no extractable text")
            
            return "\n".join(full_text), paragraphs_metadata
        
        except Exception as e:
            # Re-raise with context
            raise Exception(f"DOCX extraction failed: {str(e)}")
